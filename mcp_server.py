#!/usr/bin/env python3
"""
数学建模论文全方位评审助手 - MCP Server v1.0
基于 FastMCP 框架，提供完整的论文评审工作流

架构: 本地 MCP Server + SiliconFlow LLM API
功能: 文档解析 -> 多维度评审 -> 代码生成 -> 图表建议 -> Word 报告生成
"""

import os
import sys
import json
import tempfile
import re
from pathlib import Path
from typing import Optional, Dict, List, Any
from urllib.parse import urlparse

from fastmcp import FastMCP, Context
import yaml
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("警告: PyPDF2 未安装，PDF 解析功能不可用。请运行: pip install PyPDF2")

# 初始化配置
_root = Path(__file__).parent
if str(_root) not in sys.path:
    sys.path.insert(0, str(_root))

# 加载配置
_config_path = _root / "config.yaml"
if _config_path.exists():
    _config = yaml.safe_load(_config_path.read_text("utf-8"))
else:
    _config = {}

# 环境变量
SILICONFLOW_API_KEY = os.getenv("SILICONFLOW_API_KEY", "")
LLM_MODEL = os.getenv("LLM_MODEL", "Qwen/Qwen2.5-VL-72B-Instruct")
LLM_BASE_URL = os.getenv("LLM_BASE_URL", "https://api.siliconflow.cn/v1")

# 初始化 FastMCP
mcp = FastMCP("math_modeling_paper_reviewer")

# 数据存储目录
_data_dir = _root / "data"
_data_dir.mkdir(exist_ok=True)

# ============================================================
# MCP Tool 1: 上传并解析论文
# ============================================================
@mcp.tool()
async def upload_and_parse_paper(
    file_path: str,
    paper_type: str = "math_modeling"
) -> Dict[str, Any]:
    """
    上传数学建模竞赛论文并解析内容
    
    支持的文件格式：
    - PDF (.pdf)
    - Word (.docx)
    - 纯文本 (.txt)
    
    支持的输入方式：
    - 本地文件路径
    - HTTP/HTTPS URL
    - S3 URL (s3://...)
    
    Args:
        file_path: 论文文件路径或 URL（PDF/DOCX/TXT）
        paper_type: 论文类型，默认 math_modeling
    
    Returns:
        解析结果，包含论文元数据和全文内容
    """
    print(f"[MCP] 解析论文: {file_path}")
    
    # 检测输入类型：URL 还是本地路径
    is_url = file_path.startswith(('http://', 'https://', 's3://'))
    
    local_file_path = file_path
    temp_file = None
    
    try:
        # 如果是 URL，先下载文件
        if is_url:
            print(f"[MCP] 检测到 URL，正在下载文件...")
            local_file_path = _download_file_from_url(file_path)
            if not local_file_path:
                return {"error": f"文件下载失败: {file_path}"}
            temp_file = local_file_path  # 标记为临时文件，需要清理
        
        # 读取文件
        if not Path(local_file_path).exists():
            return {"error": f"文件不存在: {local_file_path}"}
        
        # 提取文本内容
        content = _extract_text_from_file(local_file_path)
        
        # 提取元数据
        metadata = {
            "file_name": Path(local_file_path).name,
            "file_size": Path(local_file_path).stat().st_size,
            "paper_type": paper_type,
            "word_count": len(content),
            "source": "url" if is_url else "local"
        }
        
        # 保存到临时存储
        paper_id = _save_paper_data(local_file_path, content, metadata)
        
        return {
            "success": True,
            "paper_id": paper_id,
            "metadata": metadata,
            "preview": content[:1000] + "..." if len(content) > 1000 else content,
            "message": f"成功解析论文（ID: {paper_id}），共 {metadata['word_count']} 字。请使用 paper_id='{paper_id}' 和 paper_type='{paper_type}' 进行评审。"
        }
    
    except Exception as e:
        return {"error": f"解析失败: {str(e)}"}
    
    finally:
        # 清理临时下载的文件
        if temp_file and Path(temp_file).exists():
            try:
                Path(temp_file).unlink()
                print(f"[MCP] 已清理临时文件: {temp_file}")
            except:
                pass


# ============================================================
# 论文类型评审配置
# ============================================================
REVIEW_CONFIGS = {
    "math_modeling": {
        "name": "数学建模论文",
        "aspects": [
            {"name": "问题分析", "max_score": 20, "weight": 20},
            {"name": "模型建立", "max_score": 30, "weight": 30},
            {"name": "求解方法", "max_score": 20, "weight": 20},
            {"name": "结果分析", "max_score": 15, "weight": 15},
            {"name": "论文写作", "max_score": 15, "weight": 15},
        ]
    },
    "academic": {
        "name": "学术论文",
        "aspects": [
            {"name": "创新性", "max_score": 25, "weight": 25},
            {"name": "方法论", "max_score": 25, "weight": 25},
            {"name": "实验与结果", "max_score": 20, "weight": 20},
            {"name": "文献综述", "max_score": 15, "weight": 15},
            {"name": "写作质量", "max_score": 15, "weight": 15},
        ]
    },
    "thesis": {
        "name": "学位论文",
        "aspects": [
            {"name": "研究意义", "max_score": 15, "weight": 15},
            {"name": "文献综述", "max_score": 20, "weight": 20},
            {"name": "研究方法", "max_score": 20, "weight": 20},
            {"name": "研究结果", "max_score": 25, "weight": 25},
            {"name": "论文规范", "max_score": 20, "weight": 20},
        ]
    },
    "course": {
        "name": "课程论文",
        "aspects": [
            {"name": "问题理解", "max_score": 20, "weight": 20},
            {"name": "内容质量", "max_score": 30, "weight": 30},
            {"name": "分析能力", "max_score": 20, "weight": 20},
            {"name": "结构组织", "max_score": 15, "weight": 15},
            {"name": "格式规范", "max_score": 15, "weight": 15},
        ]
    }
}

# ============================================================
# MCP Tool 2: 多维度论文评审
# ============================================================
@mcp.tool()
async def comprehensive_review(
    paper_id: str,
    review_aspects: Optional[List[str]] = None,
    paper_type: str = "math_modeling"
) -> Dict[str, Any]:
    """
    对已上传的论文进行多维度评审
    
    支持四种论文类型，每种使用不同的评审维度和分值体系：
    - math_modeling: 数学建模论文 (问题分析20 + 模型建立30 + 求解方法20 + 结果分析15 + 论文写作15)
    - academic: 学术论文 (创新性25 + 方法论25 + 实验与结果20 + 文献综述15 + 写作质量15)
    - thesis: 学位论文 (研究意义15 + 文献综述20 + 研究方法20 + 研究结果25 + 论文规范20)
    - course: 课程论文 (问题理解20 + 内容质量30 + 分析能力20 + 结构组织15 + 格式规范15)
    
    Args:
        paper_id: 论文ID
        review_aspects: 评审维度列表，默认使用对应论文类型的全部维度
        paper_type: 论文类型（math_modeling/academic/thesis/course）
    
    Returns:
        详细评审报告，包含各维度得分、总分、等级和改进建议
    """
    print(f"[MCP] 评审论文: {paper_id}, 类型: {paper_type}")
    
    # 加载论文数据
    paper_data = _load_paper_data(paper_id)
    if not paper_data:
        return {"error": "论文未找到，请先使用 upload_and_parse_paper 上传论文"}
    
    # 获取论文类型配置
    config = REVIEW_CONFIGS.get(paper_type, REVIEW_CONFIGS["math_modeling"])
    
    # 确定评审维度
    if review_aspects:
        # 用户自定义维度，使用默认分值
        aspects_to_review = [
            {"name": aspect, "max_score": 20, "weight": 20}
            for aspect in review_aspects
        ]
    else:
        aspects_to_review = config["aspects"]
    
    # 调用 LLM 进行逐维度评审
    review_results = []
    for aspect_config in aspects_to_review:
        result = _review_single_aspect(
            paper_data["content"],
            aspect_config["name"],
            aspect_config["max_score"]
        )
        review_results.append(result)
    
    # 计算加权总分
    total_max = sum(a["max_score"] for a in aspects_to_review)
    total_score = sum(r["score"] for r in review_results)
    percentage = round(total_score / total_max * 100, 2)
    
    # 确定等级
    if percentage >= 90:
        grade = "优秀"
    elif percentage >= 80:
        grade = "良好"
    elif percentage >= 70:
        grade = "中等"
    elif percentage >= 60:
        grade = "及格"
    else:
        grade = "不及格"
    
    # 生成评审报告
    report = {
        "paper_id": paper_id,
        "paper_type": config["name"],
        "total_score": total_score,
        "max_score": total_max,
        "percentage": percentage,
        "grade": grade,
        "aspects": review_results,
        "overall_suggestion": _generate_overall_suggestion(review_results)
    }
    
    # 保存评审报告
    _save_review_report(paper_id, report)
    
    return report


# ============================================================
# MCP Tool 3: 代码生成
# ============================================================
@mcp.tool()
async def generate_solution_code(
    problem_description: str,
    algorithm_type: str = "optimization",
    programming_language: str = "python"
) -> Dict[str, Any]:
    """
    根据问题描述生成求解代码
    
    Args:
        problem_description: 问题描述
        algorithm_type: 算法类型（optimization/prediction/classification/clustering）
        programming_language: 编程语言（python/matlab）
    
    Returns:
        完整的求解代码，包含详细注释
    """
    print(f"[MCP] 生成 {programming_language} 代码")
    
    # 构建提示词
    prompt = f"""
你是一个数学建模竞赛代码生成专家。

问题描述：
{problem_description}

算法类型：{algorithm_type}
编程语言：{programming_language}

请生成完整的求解代码，要求：
1. 代码必须可以直接运行
2. 包含详细的注释说明
3. 使用标准的数学建模代码结构
4. 包含数据预处理、模型建立、求解、结果输出
5. 如有需要，包含可视化代码

输出格式：
```{programming_language}
# 代码在这里
```
"""
    
    # 调用 LLM
    code = _call_llm(prompt)
    
    # 保存代码
    code_id = f"code_{hash(problem_description) % 10000:04d}"
    _save_code(code_id, code, problem_description)
    
    return {
        "code_id": code_id,
        "code": code,
        "language": programming_language,
        "message": f"已生成 {programming_language} 求解代码"
    }


# ============================================================
# MCP Tool 4: 图表生成建议
# ============================================================
@mcp.tool()
async def generate_visualization_code(
    data_description: str,
    chart_type: str = "line",
    programming_language: str = "python"
) -> Dict[str, Any]:
    """
    生成数据可视化代码
    
    Args:
        data_description: 数据描述
        chart_type: 图表类型（line/scatter/bar/pie/heatmap/3d）
        programming_language: 编程语言
    
    Returns:
        可视化代码和建议
    """
    print(f"[MCP] 生成 {chart_type} 图代码")
    
    # 构建提示词
    prompt = f"""
你是一个数据可视化专家。

数据描述：{data_description}
图表类型：{chart_type}
编程语言：{programming_language}

请生成绘制该图表的代码，要求：
1. 使用 matplotlib/seaborn（Python）或相应的 MATLAB 函数
2. 图表美观、专业，符合学术论文标准
3. 包含完整的标题、坐标轴标签、图例
4. 代码可以直接运行并生成图表
5. 添加中文注释

输出格式：
```{programming_language}
# 绘图代码
```

同时提供图表设计建议。
"""
    
    # 调用 LLM
    result = _call_llm(prompt)
    
    return {
        "code": result,
        "chart_type": chart_type,
        "language": programming_language,
        "message": f"已生成 {chart_type} 图代码"
    }


# ============================================================
# MCP Tool 5: 论文结构优化建议
# ============================================================
@mcp.tool()
async def suggest_structure_optimization(
    paper_id: str
) -> Dict[str, Any]:
    """
    分析论文结构并给出优化建议
    
    Args:
        paper_id: 论文ID
    
    Returns:
        结构分析报告和优化建议
    """
    print(f"[MCP] 分析论文结构: {paper_id}")
    
    # 加载论文
    paper_data = _load_paper_data(paper_id)
    if not paper_data:
        return {"error": "论文未找到"}
    
    # 构建提示词
    prompt = f"""
请分析以下数学建模论文的结构，并给出优化建议。

标准数学建模论文结构应该包含：
1. 摘要（Abstract）- 300-500字
2. 问题重述（Problem Restatement）
3. 模型假设（Assumptions）
4. 符号说明（Notations）
5. 模型建立与求解（Model Formulation and Solution）
6. 结果分析（Results Analysis）
7. 模型评价与推广（Model Evaluation）
8. 参考文献（References）
9. 附录（Appendix）

论文内容：
{paper_data['content'][:3000]}

请给出：
1. 当前结构分析（包含/缺失哪些部分）
2. 结构优化建议
3. 各章节字数分配建议
"""
    
    # 调用 LLM
    suggestion = _call_llm(prompt)
    
    return {
        "paper_id": paper_id,
        "structure_analysis": suggestion,
        "message": "已完成结构分析"
    }


# ============================================================
# MCP Tool 6: 生成评审报告（Word格式）
# ============================================================
@mcp.tool()
async def generate_review_report_docx(
    paper_id: str,
    output_dir: str = "output"
) -> Dict[str, Any]:
    """
    生成 Word 格式的评审报告
    
    Args:
        paper_id: 论文ID
        output_dir: 输出目录
    
    Returns:
        生成的 Word 文件路径
    """
    print(f"[MCP] 生成评审报告: {paper_id}")
    
    # 加载评审报告
    review_report = _load_review_report(paper_id)
    if not review_report:
        return {"error": "评审报告未找到，请先执行 comprehensive_review"}
    
    # 创建 Word 文档
    doc = Document()
    
    # 标题
    title = doc.add_heading('论文评审报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('一、基本信息', level=1)
    doc.add_paragraph(f'论文ID: {paper_id}')
    doc.add_paragraph(f'论文类型: {review_report.get("paper_type", "未指定")}')
    doc.add_paragraph(f'总分: {review_report["total_score"]}/{review_report["max_score"]} ({review_report["percentage"]}%)')
    doc.add_paragraph(f'等级: {review_report.get("grade", "未评定")}')
    
    # 各维度评分
    doc.add_heading('二、各维度评分', level=1)
    table = doc.add_table(rows=len(review_report["aspects"]) + 1, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '评审维度'
    hdr_cells[1].text = '得分'
    hdr_cells[2].text = '满分'
    hdr_cells[3].text = '得分率'
    
    # 数据
    for i, aspect in enumerate(review_report["aspects"], 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = aspect["aspect"]
        row_cells[1].text = str(aspect["score"])
        row_cells[2].text = str(aspect["max_score"])
        row_cells[3].text = f'{aspect["score_rate"]}%'
    
    # 详细评审
    doc.add_heading('三、详细评审意见', level=1)
    for aspect in review_report["aspects"]:
        doc.add_heading(aspect["aspect"], level=2)
        doc.add_paragraph(f'得分: {aspect["score"]}/{aspect["max_score"]}')
        doc.add_paragraph('优势:')
        for strength in aspect["strengths"]:
            doc.add_paragraph(strength, style='List Bullet')
        doc.add_paragraph('改进建议:')
        for suggestion in aspect["suggestions"]:
            doc.add_paragraph(suggestion, style='List Bullet')
    
    # 总体建议
    doc.add_heading('四、总体建议', level=1)
    doc.add_paragraph(review_report["overall_suggestion"])
    
    # 保存文档
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)
    
    file_name = f"评审报告_{paper_id}.docx"
    doc.save(output_path / file_name)
    
    return {
        "success": True,
        "file_path": str(output_path / file_name),
        "message": f"评审报告已生成: {file_name}"
    }


# ============================================================
# 辅助函数
# ============================================================

def _download_file_from_url(url: str) -> Optional[str]:
    """
    从 URL 下载文件到临时目录
    
    支持：
    - HTTP/HTTPS URL（包括 Nexent 代理 URL）
    - S3 URL (s3://bucket/key)
    - 本地文件路径（直接返回）
    
    对 Nexent 环境的特殊处理：
    - Nexent 代理 URL: http://localhost:5013/api/nb/v1/file/fetch?presigned_url=...
      自动提取内嵌的 presigned_url 并下载
    - Docker 内部 MinIO 地址 (nexent-minio:9000) 自动替换为可达地址
    - S3 URL 尝试转换为 MinIO HTTP URL 下载
    
    Returns:
        本地临时文件路径，失败返回 None
    """
    try:
        from urllib.parse import urlparse, parse_qs, unquote
        
        # 创建临时目录
        temp_dir = _data_dir / "temp_downloads"
        temp_dir.mkdir(exist_ok=True)
        
        # 生成临时文件名
        # 从 URL 中提取文件扩展名
        file_ext = '.tmp'
        if '.docx' in url.lower():
            file_ext = '.docx'
        elif '.pdf' in url.lower():
            file_ext = '.pdf'
        elif '.txt' in url.lower():
            file_ext = '.txt'
        
        temp_filename = f"download_{abs(hash(url)) % 10000:04d}{file_ext}"
        temp_path = temp_dir / temp_filename
        
        print(f"[MCP] 原始 URL: {url[:200]}..." if len(url) > 200 else f"[MCP] 原始 URL: {url}")
        
        # 处理不同类型的 URL
        download_url = url
        
        # === 1. 检测 Nexent 代理 URL 格式 ===
        # 格式: http://localhost:5013/api/nb/v1/file/fetch?presigned_url=...
        if 'presigned_url=' in url or '/api/nb/v1/file/fetch' in url:
            print(f"[MCP] 检测到 Nexent 代理 URL，正在提取真实的文件地址...")
            parsed = urlparse(url)
            query_params = parse_qs(parsed.query)
            
            if 'presigned_url' in query_params:
                # URL 解码
                presigned_url = unquote(query_params['presigned_url'][0])
                print(f"[MCP] 提取到 presigned_url: {presigned_url[:200]}..." if len(presigned_url) > 200 else f"[MCP] 提取到 presigned_url: {presigned_url}")
                download_url = presigned_url
            else:
                print(f"[MCP] 警告: Nexent 代理 URL 中未找到 presigned_url 参数")
                print(f"[MCP] 尝试直接访问代理 URL...")
                download_url = url  # 回退到原始 URL
        
        # === 2. 处理 S3 URL ===
        if download_url.startswith('s3://'):
            print(f"[MCP] 检测到 S3 URL，尝试转换为 MinIO HTTP URL...")
            # S3 URL 格式: s3://bucket/key
            # 尝试转换为 MinIO HTTP URL: http://localhost:9000/bucket/key
            s3_path = download_url[5:]  # 去掉 s3://
            slash_idx = s3_path.find('/')
            if slash_idx > 0:
                bucket = s3_path[:slash_idx]
                key = s3_path[slash_idx + 1:]
                # 尝试多种 MinIO 地址
                minio_hosts = [
                    'localhost:9000',
                    '127.0.0.1:9000',
                    '192.168.2.1:9000',
                    'host.docker.internal:9000',
                ]
                for host in minio_hosts:
                    candidate = f"http://{host}/{bucket}/{key}"
                    print(f"[MCP] 尝试 MinIO 地址: {candidate}")
                    try:
                        resp = requests.head(candidate, timeout=5)
                        if resp.status_code < 500:
                            download_url = candidate
                            print(f"[MCP] MinIO 连接成功: {candidate}")
                            break
                    except Exception:
                        continue
                else:
                    print(f"[MCP] 所有 MinIO 地址均无法连接")
                    print(f"[MCP] 提示: 请确保 MinIO 服务在本地运行且端口 9000 可访问")
                    # 最后使用 localhost 作为默认尝试
                    download_url = f"http://localhost:9000/{bucket}/{key}"
            else:
                print(f"[MCP] 无效的 S3 URL 格式: {download_url}")
                return None
        
        # === 3. Docker 内部地址替换 ===
        if 'nexent-minio:9000' in download_url:
            print(f"[MCP] 检测到 Docker 内部地址 nexent-minio:9000，尝试替换...")
            # 按优先级尝试
            replacements = [
                ('nexent-minio:9000', 'localhost:9000'),
                ('nexent-minio:9000', '127.0.0.1:9000'),
                ('nexent-minio:9000', '192.168.2.1:9000'),
                ('nexent-minio:9000', 'host.docker.internal:9000'),
            ]
            original_url = download_url
            for old, new in replacements:
                candidate = original_url.replace(old, new)
                print(f"[MCP] 尝试: {candidate[:200]}..." if len(candidate) > 200 else f"[MCP] 尝试: {candidate}")
                try:
                    resp = requests.head(candidate, timeout=5)
                    if resp.status_code < 500:
                        download_url = candidate
                        print(f"[MCP] 连接成功! 使用: {candidate[:150]}..." if len(candidate) > 150 else f"[MCP] 连接成功! 使用: {candidate}")
                        break
                except Exception:
                    continue
            else:
                # 都没连上，使用 localhost:9000 作为默认值
                print(f"[MCP] 所有替换均失败，使用默认 localhost:9000 尝试下载")
                download_url = original_url.replace('nexent-minio:9000', 'localhost:9000')
        
        # === 4. 下载文件 ===
        if download_url.startswith(('http://', 'https://')):
            print(f"[MCP] 正在下载: {download_url[:200]}..." if len(download_url) > 200 else f"[MCP] 正在下载: {download_url}")
            print(f"[MCP] 保存到: {temp_path}")
            
            # 尝试下载，带重试和多种地址回退
            last_error = None
            urls_to_try = [download_url]
            
            # 如果 URL 中包含 localhost，也尝试 127.0.0.1
            if 'localhost' in download_url:
                urls_to_try.append(download_url.replace('localhost', '127.0.0.1'))
            
            for try_url in urls_to_try:
                try:
                    response = requests.get(try_url, timeout=60, stream=True)
                    response.raise_for_status()
                    
                    # 写入文件
                    total_size = 0
                    with open(temp_path, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                                total_size += len(chunk)
                    
                    print(f"[MCP] 文件下载成功: {temp_path} ({total_size} 字节)")
                    return str(temp_path)
                except Exception as e:
                    last_error = e
                    print(f"[MCP] URL {try_url[:100]}... 下载失败: {str(e)}")
                    continue
            
            # 所有尝试都失败
            print(f"[MCP] 所有下载尝试均失败，最后错误: {last_error}")
            
            # 提供诊断信息
            print(f"[MCP] === 诊断信息 ===")
            print(f"[MCP] 1. 确认文件服务器是否在运行")
            print(f"[MCP] 2. 如果 Nexent 使用 Docker，确保 MinIO 端口已映射 (9000:9000)")
            print(f"[MCP] 3. 如果 Nexent 是桌面应用，确保 MCP Server 和 Nexent 在同一网络")
            print(f"[MCP] 4. 尝试在浏览器中直接访问该 URL 检查是否可以下载")
            return None
        
        else:
            print(f"[MCP] 不支持的 URL 格式: {download_url}")
            return None
    
    except Exception as e:
        print(f"[MCP] 文件下载失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def _call_llm(prompt: str) -> str:
    """调用 SiliconFlow LLM API"""
    headers = {
        "Authorization": f"Bearer {SILICONFLOW_API_KEY}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": LLM_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 4096,
        "temperature": 0.7
    }
    
    try:
        response = requests.post(
            f"{LLM_BASE_URL}/chat/completions",
            headers=headers,
            json=data,
            timeout=60
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"LLM 调用失败: {str(e)}"


def _extract_text_from_file(file_path: str) -> str:
    """从文件中提取文本（支持 PDF、DOCX、TXT）"""
    try:
        if file_path.endswith('.txt'):
            return Path(file_path).read_text(encoding='utf-8')
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():  # 跳过空段落
                    text_parts.append(para.text)
            return '\n'.join(text_parts)
        elif file_path.endswith('.pdf'):
            if not PDF_SUPPORT:
                return "错误: PyPDF2 未安装。请运行 'pip install PyPDF2' 来启用 PDF 解析功能。"
            
            pdf_text = []
            reader = PdfReader(file_path)
            
            # 提取每页文本
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text:
                        pdf_text.append(f"--- 第 {page_num} 页 ---\n{text}")
                except Exception as e:
                    pdf_text.append(f"--- 第 {page_num} 页 ---\n[解析失败: {str(e)}]")
            
            return '\n\n'.join(pdf_text)
        else:
            return f"不支持的文件格式: {file_path}\n支持的格式: .txt, .docx, .pdf"
    except Exception as e:
        return f"文件解析失败: {str(e)}"


def _save_paper_data(file_path: str, content: str, metadata: dict):
    """保存论文数据"""
    paper_id = f"paper_{abs(hash(file_path)) % 10000:04d}"
    data = {
        "paper_id": paper_id,
        "content": content,
        "metadata": metadata
    }
    (_data_dir / f"{paper_id}.json").write_text(
        json.dumps(data, ensure_ascii=False, indent=2),
        encoding='utf-8'
    )
    return paper_id


def _load_paper_data(paper_id: str) -> Optional[dict]:
    """加载论文数据"""
    file_path = _data_dir / f"{paper_id}.json"
    if file_path.exists():
        return json.loads(file_path.read_text(encoding='utf-8'))
    return None


def _review_single_aspect(content: str, aspect: str, max_score: int = 20) -> dict:
    """评审单个维度"""
    score_ranges = [
        (0.9, f"{int(max_score*0.9)}-{max_score}分: 优秀"),
        (0.8, f"{int(max_score*0.8)}-{int(max_score*0.9)-1}分: 良好"),
        (0.7, f"{int(max_score*0.7)}-{int(max_score*0.8)-1}分: 中等"),
        (0.6, f"{int(max_score*0.6)}-{int(max_score*0.7)-1}分: 及格"),
    ]
    score_desc = "\n".join(f"- {desc}" for _, desc in score_ranges)
    
    prompt = f"""
请评审以下论文的"{aspect}"维度。

评分标准：
{score_desc}
- 0-{int(max_score*0.6)-1}分: 不及格

论文内容：
{content[:2000]}

请给出：
1. 得分（0-{max_score}）
2. 优势（至少2点）
3. 改进建议（至少2点）

输出格式：
得分: X
优势:
1. ...
2. ...
改进建议:
1. ...
2. ...
"""
    
    result = _call_llm(prompt)
    
    # 解析结果
    score_match = re.search(r'得分[:：]\s*(\d+)', result)
    score = int(score_match.group(1)) if score_match else int(max_score * 0.75)
    score = min(score, max_score)  # 确保不超过满分
    
    return {
        "aspect": aspect,
        "score": score,
        "max_score": max_score,
        "score_rate": round(score / max_score * 100, 2),
        "strengths": re.findall(r'\d+\.\s*(.+)', result.split('改进建议')[0])[:3],
        "suggestions": re.findall(r'\d+\.\s*(.+)', result.split('改进建议')[1])[:3] if '改进建议' in result else []
    }


def _generate_overall_suggestion(aspects: list) -> str:
    """生成总体建议"""
    prompt = f"""
根据以下各维度评审结果，生成总体改进建议（200-300字）：

{json.dumps(aspects, ensure_ascii=False, indent=2)}
"""
    return _call_llm(prompt)


def _save_review_report(paper_id: str, report: dict):
    """保存评审报告"""
    (_data_dir / f"review_{paper_id}.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2),
        encoding='utf-8'
    )


def _load_review_report(paper_id: str) -> Optional[dict]:
    """加载评审报告"""
    file_path = _data_dir / f"review_{paper_id}.json"
    if file_path.exists():
        return json.loads(file_path.read_text(encoding='utf-8'))
    return None


def _save_code(code_id: str, code: str, description: str):
    """保存代码"""
    (_data_dir / f"{code_id}.py").write_text(code, encoding='utf-8')
    (_data_dir / f"{code_id}.json").write_text(
        json.dumps({"code_id": code_id, "description": description}, ensure_ascii=False, indent=2),
        encoding='utf-8'
    )


# ============================================================
# 启动 MCP 服务器
# ============================================================
if __name__ == "__main__":
    import uvicorn
    
    print("=" * 60)
    print("数学建模论文全方位评审助手 - MCP Server")
    print("=" * 60)
    print(f"API Key: {'已配置' if SILICONFLOW_API_KEY else '未配置'}")
    print(f"LLM Model: {LLM_MODEL}")
    print(f"LLM Base URL: {LLM_BASE_URL}")
    print(f"Data Directory: {_data_dir}")
    print("=" * 60)
    
    if not SILICONFLOW_API_KEY:
        print("⚠️ 警告: 未配置 SILICONFLOW_API_KEY，请在 .env 文件中设置")
    
    # 启动 FastMCP 服务器
    mcp.run(transport="sse", host="0.0.0.0", port=8004)
