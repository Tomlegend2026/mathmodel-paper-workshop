#!/usr/bin/env python3
"""
MCP 服务器测试脚本
用于验证 PDF 和 Word 文件解析功能
"""

import sys
from pathlib import Path

# 添加项目根目录到路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

def test_pdf_support():
    """测试 PDF 解析功能"""
    print("=" * 60)
    print("测试 1: PDF 解析支持")
    print("=" * 60)
    
    try:
        from PyPDF2 import PdfReader
        print("✅ PyPDF2 已安装")
        
        # 检查是否有测试文件
        test_pdf = project_root / "test_sample.pdf"
        if test_pdf.exists():
            print(f"📄 找到测试文件: {test_pdf}")
            reader = PdfReader(test_pdf)
            print(f"📊 PDF 页数: {len(reader.pages)}")
            
            # 提取第一页文本
            text = reader.pages[0].extract_text()
            print(f"📝 第一页文本长度: {len(text)} 字符")
            print(f"📋 预览: {text[:200]}...")
        else:
            print("⚠️  未找到测试文件 test_sample.pdf")
            print("💡 提示: 请放置一个 PDF 文件进行测试")
        
        return True
    except ImportError as e:
        print(f"❌ PyPDF2 未安装: {e}")
        return False
    except Exception as e:
        print(f"❌ PDF 解析失败: {e}")
        return False


def test_word_support():
    """测试 Word 解析功能"""
    print("\n" + "=" * 60)
    print("测试 2: Word (.docx) 解析支持")
    print("=" * 60)
    
    try:
        from docx import Document
        print("✅ python-docx 已安装")
        
        # 检查是否有测试文件
        test_docx = project_root / "test_sample.docx"
        if test_docx.exists():
            print(f"📄 找到测试文件: {test_docx}")
            doc = Document(test_docx)
            print(f"📊 段落数量: {len(doc.paragraphs)}")
            
            # 提取文本
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = '\n'.join(text_parts)
            print(f"📝 文本长度: {len(full_text)} 字符")
            print(f"📋 预览: {full_text[:200]}...")
        else:
            print("⚠️  未找到测试文件 test_sample.docx")
            print("💡 提示: 请放置一个 Word 文档进行测试")
        
        return True
    except ImportError as e:
        print(f"❌ python-docx 未安装: {e}")
        return False
    except Exception as e:
        print(f"❌ Word 解析失败: {e}")
        return False


def test_mcp_server():
    """测试 MCP 服务器导入"""
    print("\n" + "=" * 60)
    print("测试 3: MCP 服务器导入")
    print("=" * 60)
    
    try:
        import mcp_server
        print("✅ MCP 服务器模块导入成功")
        
        # 检查关键函数
        if hasattr(mcp_server, '_extract_text_from_file'):
            print("✅ _extract_text_from_file 函数存在")
        else:
            print("❌ _extract_text_from_file 函数不存在")
            return False
        
        # 检查 PDF 支持标志
        if hasattr(mcp_server, 'PDF_SUPPORT'):
            if mcp_server.PDF_SUPPORT:
                print("✅ PDF 支持已启用")
            else:
                print("⚠️  PDF 支持未启用（PyPDF2 未安装）")
        
        return True
    except Exception as e:
        print(f"❌ MCP 服务器导入失败: {e}")
        return False


def test_text_extraction():
    """测试文本提取函数"""
    print("\n" + "=" * 60)
    print("测试 4: 文本提取函数")
    print("=" * 60)
    
    try:
        from mcp_server import _extract_text_from_file
        
        # 创建测试 TXT 文件
        test_txt = project_root / "test_sample.txt"
        test_txt.write_text("这是一个测试文本文件。\n用于测试文本提取功能。", encoding='utf-8')
        
        print(f"📄 创建测试文件: {test_txt}")
        
        # 测试 TXT 提取
        result = _extract_text_from_file(str(test_txt))
        print(f"✅ TXT 提取成功")
        print(f"📝 提取内容: {result}")
        
        # 清理测试文件
        test_txt.unlink()
        print("🗑️  测试文件已清理")
        
        return True
    except Exception as e:
        print(f"❌ 文本提取测试失败: {e}")
        return False


def main():
    """运行所有测试"""
    print("\n" + "🧪 MCP 服务器功能测试")
    print("=" * 60)
    print(f"项目路径: {project_root}")
    print()
    
    results = []
    
    # 运行测试
    results.append(("PDF 支持", test_pdf_support()))
    results.append(("Word 支持", test_word_support()))
    results.append(("MCP 服务器", test_mcp_server()))
    results.append(("文本提取", test_text_extraction()))
    
    # 汇总结果
    print("\n" + "=" * 60)
    print("测试结果汇总")
    print("=" * 60)
    
    passed = sum(1 for _, result in results if result)
    total = len(results)
    
    for name, result in results:
        status = "✅ 通过" if result else "❌ 失败"
        print(f"{name:20s} {status}")
    
    print("-" * 60)
    print(f"总计: {passed}/{total} 测试通过")
    
    if passed == total:
        print("\n🎉 所有测试通过！MCP 服务器配置正确。")
        print("\n下一步:")
        print("1. 启动 MCP 服务器: python mcp_server.py")
        print("2. 在 AI 客户端中配置 MCP 连接")
        print("3. 开始使用论文解析功能")
    else:
        print("\n⚠️  部分测试失败，请检查错误信息并修复。")
        print("\n常见问题:")
        print("- 确保已安装所有依赖: pip install PyPDF2 python-docx")
        print("- 检查 Python 环境是否正确")
        print("- 查看上面的错误信息进行调试")
    
    return passed == total


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
